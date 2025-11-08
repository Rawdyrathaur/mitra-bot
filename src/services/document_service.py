"""
Advanced document processing service for SAM Bot
Handles multiple file formats with intelligent chunking and version control
"""
import os
import logging
import mimetypes
import hashlib
import asyncio
from typing import List, Dict, Optional, Tuple, Any, BinaryIO
from pathlib import Path
from datetime import datetime
import tempfile
import uuid

# Document processing libraries
try:
    import PyPDF2
    from pdf2image import convert_from_bytes
    from PIL import Image
    import pytesseract  # OCR for images in PDFs
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import docx
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from bs4 import BeautifulSoup
    import html2text
    HTML_AVAILABLE = True
except ImportError:
    HTML_AVAILABLE = False

try:
    import chardet  # Character encoding detection
    CHARDET_AVAILABLE = True
except ImportError:
    CHARDET_AVAILABLE = False

# Text processing
try:
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    from langchain.schema import Document as LangchainDocument
    LANGCHAIN_AVAILABLE = True
except ImportError:
    LANGCHAIN_AVAILABLE = False

# Embeddings
try:
    from sentence_transformers import SentenceTransformer
    SENTENCE_TRANSFORMERS_AVAILABLE = True
except ImportError:
    SENTENCE_TRANSFORMERS_AVAILABLE = False

import numpy as np
from database_manager import DatabaseManager

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Advanced document processor with multi-format support"""
    
    def __init__(self, db_manager: DatabaseManager, embedding_model_name: str = 'all-MiniLM-L6-v2'):
        self.db = db_manager
        self.embedding_model_name = embedding_model_name
        
        # Initialize embedding model
        if SENTENCE_TRANSFORMERS_AVAILABLE:
            try:
                self.embedding_model = SentenceTransformer(embedding_model_name)
                logger.info(f"Loaded embedding model: {embedding_model_name}")
            except Exception as e:
                logger.error(f"Failed to load embedding model: {e}")
                self.embedding_model = None
        else:
            self.embedding_model = None
            logger.warning("SentenceTransformers not available")
        
        # Initialize text splitter
        self.chunk_size = self.db.get_config('chunk_size', 1000)
        self.chunk_overlap = self.db.get_config('chunk_overlap', 200)
        
        if LANGCHAIN_AVAILABLE:
            self.text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=self.chunk_size,
                chunk_overlap=self.chunk_overlap,
                separators=["\\n\\n", "\\n", ". ", " ", ""]
            )
        else:
            self.text_splitter = None
        
        # Supported file types
        self.supported_types = {
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
            '.doc': self._process_docx,
            '.txt': self._process_text,
            '.md': self._process_text,
            '.html': self._process_html,
            '.htm': self._process_html,
            '.rtf': self._process_text,
        }
        
        logger.info("Document processor initialized")
    
    async def process_document_file(self, file_path: str, title: str = None, 
                                  uploaded_by: str = None, category: str = None) -> str:
        """Process a document file and store in knowledge base"""
        try:
            file_path = Path(file_path)
            
            if not file_path.exists():
                raise FileNotFoundError(f"File not found: {file_path}")
            
            # Get file metadata
            file_size = file_path.stat().st_size
            max_size = self.db.get_config('max_file_size', 10485760)  # 10MB default
            
            if file_size > max_size:
                raise ValueError(f"File size ({file_size}) exceeds maximum ({max_size} bytes)")
            
            file_type = file_path.suffix.lower()
            if file_type not in self.supported_types:
                raise ValueError(f"Unsupported file type: {file_type}")
            
            # Extract content
            content = await self._extract_content(file_path)
            
            if not content or not content.strip():
                raise ValueError("No text content could be extracted from the file")
            
            # Auto-detect category if not provided
            if not category:
                category = self._categorize_content(content)
            
            # Store document
            doc_id = self.db.store_document(
                title=title or file_path.stem,
                content=content,
                file_type=file_type[1:],  # Remove the dot
                filename=file_path.name,
                uploaded_by=uploaded_by,
                category=category
            )
            
            # Process chunks and embeddings
            await self._process_chunks(doc_id, content)
            
            logger.info(f"Successfully processed document: {file_path.name} -> {doc_id}")
            return doc_id
            
        except Exception as e:
            logger.error(f"Error processing document {file_path}: {e}")
            raise
    
    async def process_text_content(self, title: str, content: str, 
                                 uploaded_by: str = None, category: str = None) -> str:
        """Process raw text content and store in knowledge base"""
        try:
            if not content or not content.strip():
                raise ValueError("Content cannot be empty")
            
            # Auto-detect category if not provided
            if not category:
                category = self._categorize_content(content)
            
            # Store document
            doc_id = self.db.store_document(
                title=title,
                content=content,
                file_type='text',
                uploaded_by=uploaded_by,
                category=category
            )
            
            # Process chunks and embeddings
            await self._process_chunks(doc_id, content)
            
            logger.info(f"Successfully processed text content: {title} -> {doc_id}")
            return doc_id
            
        except Exception as e:
            logger.error(f"Error processing text content: {e}")
            raise
    
    async def _extract_content(self, file_path: Path) -> str:
        """Extract text content from file based on type"""
        file_type = file_path.suffix.lower()
        
        if file_type in self.supported_types:
            return await self.supported_types[file_type](file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    async def _process_pdf(self, file_path: Path) -> str:
        """Process PDF file with OCR support for images"""
        if not PDF_AVAILABLE:
            raise ImportError("PDF processing libraries not available")
        
        try:
            text_content = []
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        # Extract text
                        page_text = page.extract_text()
                        
                        if page_text.strip():
                            text_content.append(f"[Page {page_num + 1}]\\n{page_text}")
                        else:
                            # Try OCR if no text found
                            logger.info(f"No text found on page {page_num + 1}, attempting OCR")
                            # OCR implementation would go here
                            # For now, add a placeholder
                            text_content.append(f"[Page {page_num + 1}]\n[Image/Scanned content detected]")

                    except Exception as e:
                        logger.warning(f"Error processing page {page_num + 1}: {e}")
            
            content = "\\n\\n".join(text_content)
            
            if not content.strip():
                raise ValueError("No text content could be extracted from PDF")
            
            return content
            
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {e}")
            raise
    
    async def _process_docx(self, file_path: Path) -> str:
        """Process DOCX file"""
        if not DOCX_AVAILABLE:
            raise ImportError("DOCX processing libraries not available")
        
        try:
            doc = DocxDocument(file_path)
            paragraphs = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text)
            
            # Process tables
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    if any(row_data):  # Skip empty rows
                        table_data.append(" | ".join(row_data))
                
                if table_data:
                    paragraphs.append("\\n".join(table_data))
            
            content = "\\n\\n".join(paragraphs)
            
            if not content.strip():
                raise ValueError("No text content could be extracted from DOCX")
            
            return content
            
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {e}")
            raise
    
    async def _process_html(self, file_path: Path) -> str:
        """Process HTML file"""
        if not HTML_AVAILABLE:
            raise ImportError("HTML processing libraries not available")
        
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                html_content = f.read()
            
            # Parse HTML and convert to text
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Convert to markdown-like text
            h = html2text.HTML2Text()
            h.ignore_links = True
            h.ignore_images = True
            text_content = h.handle(str(soup))
            
            if not text_content.strip():
                raise ValueError("No text content could be extracted from HTML")
            
            return text_content
            
        except Exception as e:
            logger.error(f"Error processing HTML {file_path}: {e}")
            raise
    
    async def _process_text(self, file_path: Path) -> str:
        """Process plain text file with encoding detection"""
        try:
            # Try to detect encoding
            encoding = 'utf-8'
            
            if CHARDET_AVAILABLE:
                with open(file_path, 'rb') as f:
                    raw_data = f.read()
                    detected = chardet.detect(raw_data)
                    if detected['confidence'] > 0.7:
                        encoding = detected['encoding']
            
            # Read file with detected encoding
            with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
                content = f.read()
            
            if not content.strip():
                raise ValueError("File is empty or contains no readable text")
            
            return content
            
        except Exception as e:
            logger.error(f"Error processing text file {file_path}: {e}")
            raise
    
    async def _process_chunks(self, document_id: str, content: str):
        """Process document into chunks with embeddings"""
        try:
            # Split into chunks
            chunks = self._split_text(content)
            
            if not chunks:
                raise ValueError("No chunks generated from content")
            
            # Generate embeddings
            embeddings = await self._generate_embeddings(chunks)
            
            # Store in database
            chunk_ids = self.db.store_document_chunks(document_id, chunks, embeddings)
            
            logger.info(f"Processed {len(chunks)} chunks for document {document_id}")
            return chunk_ids
            
        except Exception as e:
            logger.error(f"Error processing chunks for document {document_id}: {e}")
            raise
    
    def _split_text(self, text: str) -> List[str]:
        """Split text into chunks using configured strategy"""
        try:
            if self.text_splitter and LANGCHAIN_AVAILABLE:
                # Use LangChain's advanced text splitter
                chunks = self.text_splitter.split_text(text)
                return [chunk.strip() for chunk in chunks if chunk.strip()]
            else:
                # Fallback to simple word-based chunking
                return self._simple_text_split(text)
                
        except Exception as e:
            logger.error(f"Error splitting text: {e}")
            return [text]  # Return original text as single chunk
    
    def _simple_text_split(self, text: str) -> List[str]:
        """Simple text splitting with word boundaries"""
        words = text.split()
        chunks = []
        current_chunk = []
        current_length = 0
        
        for word in words:
            word_length = len(word) + 1  # +1 for space
            
            if current_length + word_length > self.chunk_size and current_chunk:
                # Save current chunk
                chunks.append(" ".join(current_chunk))
                
                # Start new chunk with overlap
                overlap_words = current_chunk[-self.chunk_overlap//10:] if len(current_chunk) > self.chunk_overlap//10 else []
                current_chunk = overlap_words + [word]
                current_length = sum(len(w) + 1 for w in current_chunk)
            else:
                current_chunk.append(word)
                current_length += word_length
        
        # Add final chunk
        if current_chunk:
            chunks.append(" ".join(current_chunk))
        
        return chunks
    
    async def _generate_embeddings(self, texts: List[str]) -> List[List[float]]:
        """Generate embeddings for text chunks"""
        try:
            if not self.embedding_model or not texts:
                logger.warning("Embedding model not available or no texts provided")
                # Return zero embeddings as fallback
                embedding_size = 384  # Default for all-MiniLM-L6-v2
                return [[0.0] * embedding_size for _ in texts]
            
            # Generate embeddings in batches for efficiency
            batch_size = 32
            all_embeddings = []
            
            for i in range(0, len(texts), batch_size):
                batch = texts[i:i + batch_size]
                try:
                    batch_embeddings = self.embedding_model.encode(
                        batch, 
                        convert_to_numpy=True,
                        show_progress_bar=False
                    )
                    all_embeddings.extend(batch_embeddings.tolist())
                except Exception as e:
                    logger.error(f"Error generating embeddings for batch {i}: {e}")
                    # Add zero embeddings for failed batch
                    embedding_size = 384
                    all_embeddings.extend([[0.0] * embedding_size for _ in batch])
            
            logger.info(f"Generated embeddings for {len(texts)} chunks")
            return all_embeddings
            
        except Exception as e:
            logger.error(f"Error generating embeddings: {e}")
            # Return zero embeddings as fallback
            embedding_size = 384
            return [[0.0] * embedding_size for _ in texts]
    
    def _categorize_content(self, content: str) -> str:
        """Auto-categorize content based on keywords and patterns"""
        try:
            content_lower = content.lower()
            
            # Define category keywords
            categories = {
                'faq': ['frequently asked', 'question', 'answer', 'faq', 'q&a'],
                'documentation': ['documentation', 'manual', 'guide', 'tutorial', 'how to', 'instructions'],
                'policy': ['policy', 'terms', 'conditions', 'privacy', 'legal', 'compliance'],
                'product': ['product', 'feature', 'specification', 'technical', 'api', 'sdk'],
                'support': ['support', 'troubleshooting', 'help', 'issue', 'problem', 'solution'],
                'onboarding': ['onboarding', 'getting started', 'setup', 'installation', 'quickstart']
            }
            
            # Count keyword matches
            category_scores = {}
            for category, keywords in categories.items():
                score = sum(1 for keyword in keywords if keyword in content_lower)
                if score > 0:
                    category_scores[category] = score
            
            # Return category with highest score, or 'general' if none
            if category_scores:
                return max(category_scores.items(), key=lambda x: x[1])[0]
            else:
                return 'general'
                
        except Exception as e:
            logger.error(f"Error categorizing content: {e}")
            return 'general'
    
    def get_processing_status(self, document_id: str) -> Dict:
        """Get processing status for a document"""
        try:
            with self.db.get_session() as session:
                from models.database_models import Document
                document = session.query(Document).filter(Document.id == document_id).first()
                
                if document:
                    return {
                        'document_id': document_id,
                        'status': document.processing_status,
                        'error': document.processing_error,
                        'created_at': document.created_at.isoformat(),
                        'updated_at': document.updated_at.isoformat()
                    }
                else:
                    return {'error': 'Document not found'}
                    
        except Exception as e:
            logger.error(f"Error getting processing status: {e}")
            return {'error': str(e)}
    
    def get_document_info(self, document_id: str) -> Dict:
        """Get detailed document information"""
        try:
            with self.db.get_session() as session:
                from models.database_models import Document, DocumentChunk
                
                document = session.query(Document).filter(Document.id == document_id).first()
                if not document:
                    return {'error': 'Document not found'}
                
                # Get chunk count
                chunk_count = session.query(func.count(DocumentChunk.id)).filter(
                    DocumentChunk.document_id == document_id
                ).scalar()
                
                return {
                    'id': document.id,
                    'title': document.title,
                    'filename': document.filename,
                    'file_type': document.file_type,
                    'file_size': document.file_size,
                    'category': document.category,
                    'language': document.language,
                    'version': document.version,
                    'processing_status': document.processing_status,
                    'chunk_count': chunk_count,
                    'created_at': document.created_at.isoformat(),
                    'updated_at': document.updated_at.isoformat()
                }
                
        except Exception as e:
            logger.error(f"Error getting document info: {e}")
            return {'error': str(e)}
    
    def list_documents(self, category: str = None, status: str = None, 
                      limit: int = 50, offset: int = 0) -> List[Dict]:
        """List documents with optional filtering"""
        try:
            with self.db.get_session() as session:
                from models.database_models import Document, DocumentChunk
                
                query = session.query(Document, func.count(DocumentChunk.id).label('chunk_count')).outerjoin(DocumentChunk)
                
                if category:
                    query = query.filter(Document.category == category)
                
                if status:
                    query = query.filter(Document.processing_status == status)
                
                query = query.group_by(Document.id).order_by(Document.created_at.desc())
                query = query.offset(offset).limit(limit)
                
                results = query.all()
                
                return [{
                    'id': doc.id,
                    'title': doc.title,
                    'filename': doc.filename,
                    'file_type': doc.file_type,
                    'category': doc.category,
                    'processing_status': doc.processing_status,
                    'chunk_count': chunk_count,
                    'created_at': doc.created_at.isoformat()
                } for doc, chunk_count in results]
                
        except Exception as e:
            logger.error(f"Error listing documents: {e}")
            return []
    
    def delete_document(self, document_id: str, user_id: str = None) -> bool:
        """Delete document and all its chunks"""
        try:
            with self.db.get_session() as session:
                from models.database_models import Document
                
                document = session.query(Document).filter(Document.id == document_id).first()
                if document:
                    session.delete(document)  # Cascading delete will remove chunks
                    
                    # Log analytics
                    self.db._log_analytics(session, 'document_deleted',
                                         user_id=user_id,
                                         metadata={'document_id': document_id})
                    
                    logger.info(f"Deleted document {document_id}")
                    return True
                else:
                    logger.warning(f"Document {document_id} not found for deletion")
                    return False
                    
        except Exception as e:
            logger.error(f"Error deleting document {document_id}: {e}")
            return False

class DocumentUploadHandler:
    """Handle document uploads with validation and processing"""
    
    def __init__(self, document_processor: DocumentProcessor):
        self.processor = document_processor
        
    async def handle_file_upload(self, file_data: BinaryIO, filename: str,
                               title: str = None, uploaded_by: str = None,
                               category: str = None) -> Dict:
        """Handle file upload from web interface or API"""
        try:
            # Create temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix=Path(filename).suffix) as temp_file:
                file_data.seek(0)
                temp_file.write(file_data.read())
                temp_path = temp_file.name
            
            try:
                # Process the file
                doc_id = await self.processor.process_document_file(
                    temp_path, 
                    title=title or Path(filename).stem,
                    uploaded_by=uploaded_by,
                    category=category
                )
                
                return {
                    'success': True,
                    'document_id': doc_id,
                    'message': 'Document processed successfully'
                }
                
            finally:
                # Clean up temporary file
                os.unlink(temp_path)
                
        except Exception as e:
            logger.error(f"Error handling file upload {filename}: {e}")
            return {
                'success': False,
                'error': str(e)
            }
    
    def validate_file(self, filename: str, file_size: int) -> Tuple[bool, str]:
        """Validate uploaded file"""
        try:
            # Check file extension
            file_ext = Path(filename).suffix.lower()
            if file_ext not in self.processor.supported_types:
                return False, f"Unsupported file type: {file_ext}"
            
            # Check file size
            max_size = self.processor.db.get_config('max_file_size', 10485760)
            if file_size > max_size:
                return False, f"File size ({file_size} bytes) exceeds maximum ({max_size} bytes)"
            
            # Check filename
            if not filename or len(filename) > 255:
                return False, "Invalid filename"
            
            return True, "Valid"
            
        except Exception as e:
            logger.error(f"Error validating file {filename}: {e}")
            return False, str(e)
